"""
Document Ingestion Lambda Handler
Επεξεργάζεται έγγραφα: εξαγωγή κειμένου, chunking, αποστολή για embedding.

Serverless RAG Project - MSc Thesis
"""

import json
import os
import re
import time
import hashlib
import urllib.parse
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
import boto3

# ============================================================================
# Configuration
# ============================================================================

S3_BUCKET = os.environ.get("DOCUMENTS_BUCKET", "rag-documents")
SQS_QUEUE_URL = os.environ.get("EMBEDDING_QUEUE_URL", "")
DYNAMODB_TABLE = os.environ.get("METADATA_TABLE", "rag-metadata")
CHUNK_SIZE = int(os.environ.get("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.environ.get("CHUNK_OVERLAP", "200"))

# AWS Clients
s3 = boto3.client("s3")
sqs = boto3.client("sqs")
dynamodb = boto3.resource("dynamodb")


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class DocumentMetadata:
    """Metadata για ένα έγγραφο"""
    document_id: str
    filename: str
    file_type: str
    file_size: int
    s3_key: str
    chunk_count: int
    status: str  # processing | completed | failed
    created_at: float
    error: Optional[str] = None


@dataclass
class TextChunk:
    """Ένα κομμάτι κειμένου"""
    text: str
    chunk_index: int
    start_char: int
    end_char: int


# ============================================================================
# Text Extraction
# ============================================================================

def extract_text_pdf(content: bytes) -> str:
    """Εξαγωγή κειμένου από PDF"""
    try:
        from pypdf import PdfReader
        import io

        reader = PdfReader(io.BytesIO(content))
        pages = []

        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Σελίδα {i}]\n{text}")

        return "\n\n".join(pages)
    except ImportError:
        # Fallback αν δεν υπάρχει pypdf
        return "[PDF extraction requires pypdf library]"
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")


def extract_text_docx(content: bytes) -> str:
    """Εξαγωγή κειμένου από Word"""
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Εξαγωγή από πίνακες
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)
    except ImportError:
        return "[DOCX extraction requires python-docx library]"
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}")


def extract_text_plain(content: bytes) -> str:
    """Εξαγωγή από plain text"""
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file")


def extract_text(content: bytes, file_type: str) -> str:
    """Router για εξαγωγή κειμένου βάσει τύπου αρχείου"""
    extractors = {
        "pdf": extract_text_pdf,
        "docx": extract_text_docx,
        "doc": extract_text_docx,
        "txt": extract_text_plain,
        "md": extract_text_plain,
    }

    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor(content)


# ============================================================================
# Text Processing
# ============================================================================

def clean_text(text: str) -> str:
    """Καθαρισμός και κανονικοποίηση κειμένου"""
    # Αφαίρεση control characters
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

    # Κανονικοποίηση whitespace
    text = re.sub(r"\s+", " ", text)

    # Κανονικοποίηση quotes
    text = text.replace(""", '"').replace(""", '"')
    text = text.replace("'", "'").replace("'", "'")

    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE,
               overlap: int = CHUNK_OVERLAP) -> List[TextChunk]:
    """
    Διαίρεση κειμένου σε overlapping chunks.
    Χρησιμοποιεί sentence boundaries όπου είναι δυνατόν.
    """
    # Split σε προτάσεις
    sentence_pattern = re.compile(r"(?<=[.!?;])\s+")
    sentences = sentence_pattern.split(text)

    chunks = []
    current_text = ""
    current_start = 0
    char_pos = 0

    for sentence in sentences:
        # Αν το chunk γίνεται πολύ μεγάλο
        if len(current_text) + len(sentence) > chunk_size and current_text:
            chunks.append(TextChunk(
                text=current_text.strip(),
                chunk_index=len(chunks),
                start_char=current_start,
                end_char=char_pos
            ))

            # Νέο chunk με overlap
            overlap_start = max(0, len(current_text) - overlap)
            current_text = current_text[overlap_start:] + " " + sentence
            current_start = char_pos - (len(current_text) - len(sentence) - 1)
        else:
            current_text += (" " + sentence) if current_text else sentence

        char_pos += len(sentence) + 1

    # Τελευταίο chunk
    if current_text.strip():
        chunks.append(TextChunk(
            text=current_text.strip(),
            chunk_index=len(chunks),
            start_char=current_start,
            end_char=char_pos
        ))

    return chunks


# ============================================================================
# AWS Operations
# ============================================================================

def generate_doc_id(bucket: str, key: str) -> str:
    """Δημιουργία unique document ID"""
    content = f"{bucket}:{key}:{time.time()}"
    return hashlib.sha256(content.encode()).hexdigest()[:16]


def get_file_type(key: str) -> str:
    """Εξαγωγή file type από S3 key"""
    parts = key.rsplit(".", 1)
    return parts[1].lower() if len(parts) > 1 else "txt"


def download_from_s3(bucket: str, key: str) -> Tuple[bytes, int]:
    """Download έγγραφο από S3"""
    response = s3.get_object(Bucket=bucket, Key=key)
    content = response["Body"].read()
    return content, len(content)


def save_metadata(metadata: DocumentMetadata):
    """Αποθήκευση metadata στο DynamoDB"""
    table = dynamodb.Table(DYNAMODB_TABLE)
    item = asdict(metadata)
    item = {k: v for k, v in item.items() if v is not None}
    table.put_item(Item=item)


def send_to_queue(document_id: str, chunks: List[TextChunk],
                  metadata: Dict[str, Any]) -> int:
    """Αποστολή chunks στο SQS queue για embedding"""
    if not SQS_QUEUE_URL:
        print("Warning: SQS_QUEUE_URL not configured")
        return 0

    message = {
        "document_id": document_id,
        "chunks": [
            {"text": c.text, "chunk_index": c.chunk_index}
            for c in chunks
        ],
        "metadata": metadata
    }

    sqs.send_message(
        QueueUrl=SQS_QUEUE_URL,
        MessageBody=json.dumps(message)
    )

    return 1


# ============================================================================
# Main Processing
# ============================================================================

def process_document(bucket: str, key: str) -> DocumentMetadata:
    """
    Επεξεργασία ενός εγγράφου:
    1. Download από S3
    2. Εξαγωγή κειμένου
    3. Chunking
    4. Αποστολή για embedding
    """
    doc_id = generate_doc_id(bucket, key)
    filename = key.split("/")[-1]
    file_type = get_file_type(key)

    print(f"Processing: {doc_id} ({filename})")

    try:
        # Download
        content, file_size = download_from_s3(bucket, key)

        # Extract text
        raw_text = extract_text(content, file_type)
        clean = clean_text(raw_text)

        if not clean:
            raise ValueError("No text content extracted")

        # Chunk
        chunks = chunk_text(clean)
        print(f"Created {len(chunks)} chunks")

        # Queue for embedding
        doc_metadata = {
            "filename": filename,
            "file_type": file_type,
            "s3_key": key
        }
        send_to_queue(doc_id, chunks, doc_metadata)

        # Save metadata
        metadata = DocumentMetadata(
            document_id=doc_id,
            filename=filename,
            file_type=file_type,
            file_size=file_size,
            s3_key=key,
            chunk_count=len(chunks),
            status="processing",
            created_at=time.time()
        )
        save_metadata(metadata)

        return metadata

    except Exception as e:
        print(f"Error processing {key}: {e}")
        metadata = DocumentMetadata(
            document_id=doc_id,
            filename=filename,
            file_type=file_type,
            file_size=0,
            s3_key=key,
            chunk_count=0,
            status="failed",
            created_at=time.time(),
            error=str(e)
        )
        save_metadata(metadata)
        raise


# ============================================================================
# Lambda Handler
# ============================================================================

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Lambda entry point.

    Triggers:
    - S3 Event: Νέο έγγραφο uploaded
    - API Gateway: Manual upload
    - Direct invocation
    """
    results = []

    # S3 Event Trigger
    if "Records" in event:
        for record in event["Records"]:
            if record.get("eventSource") == "aws:s3":
                bucket = record["s3"]["bucket"]["name"]
                key = urllib.parse.unquote_plus(record["s3"]["object"]["key"])

                try:
                    metadata = process_document(bucket, key)
                    results.append({
                        "document_id": metadata.document_id,
                        "status": metadata.status,
                        "chunks": metadata.chunk_count
                    })
                except Exception as e:
                    results.append({
                        "document_id": generate_doc_id(bucket, key),
                        "status": "failed",
                        "error": str(e)
                    })

    # API / Direct invocation
    elif "body" in event or "bucket" in event:
        body = event.get("body", event)
        if isinstance(body, str):
            body = json.loads(body)

        bucket = body.get("bucket", S3_BUCKET)
        key = body.get("key")

        if not key:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Missing 'key' parameter"})
            }

        try:
            metadata = process_document(bucket, key)
            results.append({
                "document_id": metadata.document_id,
                "status": metadata.status,
                "chunks": metadata.chunk_count
            })
        except Exception as e:
            return {
                "statusCode": 500,
                "body": json.dumps({"error": str(e)})
            }

    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*"
        },
        "body": json.dumps({
            "message": f"Processed {len(results)} document(s)",
            "results": results
        })
    }


